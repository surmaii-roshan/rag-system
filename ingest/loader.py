"""
ingest/loader.py — Multi-format document loader.

Supported formats: PDF, DOCX, HTML, TXT, MD.
Returns a list of Document objects with source and page metadata.
"""

from pathlib import Path
from typing import List

import pymupdf                          # pip: pymupdf>=1.25.0  |  import: pymupdf (not fitz)
from docx import Document as DocxDoc
from bs4 import BeautifulSoup

from ingest.models import Document
from utils.logger import get_logger

log = get_logger(__name__)

# Formats this loader supports
SUPPORTED_EXTENSIONS = {".pdf", ".txt", ".md", ".docx", ".html", ".htm"}


def load_document(path: str | Path) -> List[Document]:
    """
    Load a document from disk and return a list of Document objects.
    PDFs produce one Document per page. All others produce one Document total.

    Args:
        path: Path to the file to load.

    Returns:
        List of Document objects, or empty list if loading fails.
    """
    path = Path(path)

    if not path.exists():
        log.error(f"File not found: {path}")
        return []

    ext = path.suffix.lower()

    if ext not in SUPPORTED_EXTENSIONS:
        log.warning(f"Unsupported format '{ext}': {path.name}")
        return []

    try:
        if ext == ".pdf":
            return _load_pdf(path)
        elif ext in (".txt", ".md"):
            return _load_text(path)
        elif ext == ".docx":
            return _load_docx(path)
        elif ext in (".html", ".htm"):
            return _load_html(path)
    except Exception as e:
        log.error(f"Failed to load {path.name}: {e}")
        return []

    return []


def _load_pdf(path: Path) -> List[Document]:
    """One Document per page. Skips blank pages."""
    docs = []
    pdf = pymupdf.open(str(path))
    for page_num, page in enumerate(pdf, start=1):
        text = page.get_text().strip()
        if text:
            docs.append(Document(
                content=text,
                source_file=path.name,
                page_number=page_num,
            ))
    pdf.close()
    log.debug(f"PDF '{path.name}': {len(docs)} non-empty pages extracted")
    return docs


def _load_text(path: Path) -> List[Document]:
    """Plain text or Markdown — one Document for the whole file."""
    text = path.read_text(encoding="utf-8", errors="ignore").strip()
    if not text:
        log.warning(f"Empty file: {path.name}")
        return []
    return [Document(content=text, source_file=path.name, page_number=1)]


def _load_docx(path: Path) -> List[Document]:
    """DOCX — join all non-empty paragraphs into one Document."""
    doc = DocxDoc(str(path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    text = "\n\n".join(paragraphs)
    if not text:
        log.warning(f"No text content in DOCX: {path.name}")
        return []
    return [Document(content=text, source_file=path.name, page_number=1)]


def _load_html(path: Path) -> List[Document]:
    """HTML — strip tags, return visible text as one Document."""
    html = path.read_text(encoding="utf-8", errors="ignore")
    soup = BeautifulSoup(html, "lxml")
    # Remove script and style elements
    for tag in soup(["script", "style", "meta", "link"]):
        tag.decompose()
    text = soup.get_text(separator="\n").strip()
    if not text:
        log.warning(f"No visible text in HTML: {path.name}")
        return []
    return [Document(content=text, source_file=path.name, page_number=1)]